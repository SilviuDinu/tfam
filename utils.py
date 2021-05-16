import os
import shutil
import csv
import time
import numpy as np
import math
import sys
import cv2
from PIL import Image,ImageFilter
from matplotlib.pyplot import *
import matplotlib.pyplot as plt
from os.path import abspath
from prettytable import PrettyTable
import docx
from docx import Document
from docx.shared import Inches
import pandas as pd

curr = os.path.dirname(__file__)

def build_word_doc(rows, headers, doc_name):
    df = pd.DataFrame(rows, columns=headers)
    if not os.path.exists(os.path.join(curr, doc_name)):
        document = Document()
        document.save(os.path.join(curr, doc_name))

    # open an existing document
    doc = docx.Document(os.path.join(curr, doc_name))
    t = doc.add_table(df.shape[0]+1, df.shape[1])

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            value = str(df.values[i,j])
            if os.path.isfile(value) or os.path.isdir(value):
                paragraph = t.cell(i+1,j).paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(value, width=Inches(0.65), height=Inches(0.65))
            else:
                t.cell(i+1,j).text = value

    # save the doc
    doc.save(os.path.join(curr, doc_name))

def build_graph(names, medie, r, g, b, title):       
    X = np.arange(len(names))
    fig = plt.figure(figsize=(15, 8), dpi=80)

    # ax = fig.add_axes([0,0,1,1])
    ax = fig.add_subplot(111)
    ax.bar(X + 0.00, r, color = 'r', width = 0.20, label='canal R')
    ax.bar(X + 0.20, g, color = 'g', width = 0.20, label='canal G')
    ax.bar(X + 0.40, b, color = 'b', width = 0.20, label='canal B')
    ax.bar(X + 0.60, medie, color = 'pink', width = 0.20, label='Media')

    ax.set_title(title)
    ax.legend()

    plt.xticks([i + 0.35 for i in range(len(names))], names)
    plt.title(title)
    plt.xlabel(title)
    plt.ylabel('Values')
    plt.show()

def build_graphs(data, title):
    benigne = {
        "overall": [],
        "name": [],
        "medie": [],
        "r": [],
        "g": [],
        "b": []
    }
    maligne = {
        "overall": [],
        "name": [],
        "medie": [],
        "r": [],
        "g": [],
        "b": []
    }
    for i in data:
        name, values = i
        if '_b' in name:
            benigne["overall"].append([name, values["r"], values["g"], values["b"], values["medie"]])
            benigne["name"].append(name)
            benigne["r"].append(values["r"])
            benigne["g"].append(values["g"])
            benigne["b"].append(values["b"])
            benigne["medie"].append(values["medie"])
        elif '_m' in name:
            maligne["overall"].append([name, values["r"], values["g"], values["b"], values["medie"]])
            maligne["name"].append(name)
            maligne["r"].append(values["r"])
            maligne["g"].append(values["g"])
            maligne["b"].append(values["b"])
            maligne["medie"].append(values["medie"])
    
    build_graph(benigne["name"], benigne["medie"], benigne["r"], benigne["g"], benigne["b"], title)
    build_graph(maligne["name"], maligne["medie"], maligne["r"], maligne["g"], maligne["b"], title)



def build_scaled(imgFilePath, newImgPath):
    size = 128, 128
    print('Copying images to new location...')

    for root, dirs, files in os.walk(imgFilePath):
        for idx, file in enumerate(files):
            fileName = file.split('.')[0]
            sys.stdout.flush()
            if not os.path.exists(newImgPath):
                os.makedirs(newImgPath)
            shutil.copy(os.path.join(root,file), newImgPath)
            sys.stdout.write("\bCurrent progress: %s %%\r" % (str(math.ceil(idx / len(files) * 100))))
            sys.stdout.flush()
    sys.stdout.write("\b\nDone\n")

    print('Scaling test images to %s... and deleting old ones' % (str(size)))
    for root, dirs, files in os.walk(newImgPath):
        for idx, file in enumerate(files):
            outfile = 'ISIC_' + file if not 'ISIC_' in file else file
            try:
                im = Image.open(os.path.join(newImgPath, file))
                im = im.resize(size, Image.ANTIALIAS)
                sys.stdout.flush()
                sys.stdout.write("\bCurrent progress: %s %%\r" % (str(math.ceil(idx / len(files) * 100))))
                im.save(os.path.join(newImgPath, outfile), "JPEG" ,quality=100)
                os.remove(os.path.join(newImgPath, file))
            except IOError:
                print("cannot create thumbnail for '%s'" % file)
    sys.stdout.write("\b\nDone\n")


def crop(imgFilePath, newImgPath):
    print('Cropping images...')
    for root, dirs, files in os.walk(imgFilePath):
        for idx, file in enumerate(files):

            if not os.path.exists(newImgPath):
                os.makedirs(newImgPath)
            # if idx == 0:

            new_width = 75
            new_height = 75
            try:
                im = Image.open(os.path.join(imgFilePath, file))
                width, height = im.size
                left = int(np.floor((width - new_width)/2))
                top =  int(np.floor((height - new_height)/2))
                right = int(np.floor((width + new_width)/2))
                bottom = int(np.floor((height + new_height)/2))
                
                # Cropped image of above dimension
                # (It will not change orginal image)
                im1 = im.crop((left, top, right, bottom))
                sys.stdout.flush()
                sys.stdout.write("\bCurrent progress: %s %%\r" % (str(math.ceil(idx / len(files) * 100))))
                # Shows the image in image viewer
                im1.save(os.path.join(newImgPath, file))
            except IOError:
                print("cannot crop '%s'" % file)
            
    sys.stdout.write("\bDone\n")



def dbc(img,s):
    (width, height) = img.size
    # check width == height
    assert(width == height)
    pixel = img.load()
    M = width
    # grid size must be bigger than 2 and least than M/2
    G = 256
    assert(s >= 2)
    assert(s <= M//2)
    ngrid = math.ceil(M / s)
    h = G*(s / M) # box height
    nr = np.zeros((ngrid,ngrid), dtype='int32')
    
    for i in range(ngrid):
        for j in range(ngrid):
            maxg = 0
            ming = 255
            for k in range(i*s, min((i+1)*s, M)):
                for l in range(j*s, min((j+1)*s, M)):
                    if pixel[k, l] > maxg:
                        maxg = pixel[k, l]

                    if pixel[k, l] < ming:
                        ming = pixel[k, l]
                        
            nr[i,j] = math.ceil(maxg/h) - math.ceil(ming/h) + 1

    Ns = 0
    N = np.size(nr)
    for i in range(ngrid):
        for j in range(ngrid):
            Ns += nr[i, j]
    return Ns, N, nr

def rgb(imgFilePath, newImgPath):
    first_headers = ['Image_name', 'Image','Canal', 'Ns', 'procent_box', 'fractal_dim', 'lacunaritate_m']
    lacunaritate_header = ['Image_name', 'Canal', '2', '3', '4', '5', '6', '7', '9', '10', '12', '14', '17', '20', '23', '27', '31', '37']
    original_images_header = ['Nume', 'Original', 'R', 'G', 'B']
    lacunaritate_tabel_terminal_header = ['Image_name', 'Canal']
    dim_fractala_medie_header = ['Image_name', 'Image', 'dim_fractala_medie', 'lacunaritate_medie']

    first_rows = []
    original_images_rows = []
    lacunaritate_rows = []
    dim_fractala_medie_rows = []
#  str(key) + ' => ' + str(round(lacunaritate[key], 5))

    lacunaritate_tabel_terminal = PrettyTable(lacunaritate_header)

    original_images_path = os.path.join(curr, r'data/image-data/imagini')

    print('Splitting into r, g, b then calculate lacunarity and fractal dimension...')
    b_path = os.path.join(newImgPath, r'b')
    g_path = os.path.join(newImgPath, r'g')
    r_path = os.path.join(newImgPath, r'r')

    if not os.path.exists(newImgPath):
        os.makedirs(newImgPath)

    if not os.path.exists(b_path):
        os.makedirs(b_path)
    if not os.path.exists(g_path):
        os.makedirs(g_path)
    if not os.path.exists(r_path):
        os.makedirs(r_path)

    dim_fractal_graph = []
    lac_graph = []

    for root, dirs, files in os.walk(imgFilePath):
        for idx, file in enumerate(files):

            # if idx == 0:
            img = cv2.imread(os.path.join(imgFilePath, file))
            r, g, b = cv2.split(img)

            cv2.imwrite(os.path.join(b_path, file), b)
            cv2.imwrite(os.path.join(g_path, file), g)
            cv2.imwrite(os.path.join(r_path, file), r)

            filename = file.split('.')[0]
            original_filename = file.split('ISIC_')[1]
            original_images_rows.append([filename, os.path.join(original_images_path, original_filename), os.path.join(r_path, file), os.path.join(g_path, file), os.path.join(b_path, file)])
            dim_medii = {}
            lac_medii = {}

            lacunaritate_medie = 0
            lacunaritate_medie_list = []
            dimensiune_fractala_medie_list = []

            lacunaritate_R = [filename, 'R']
            Ns, procent_box, factor_divizare, fractal_dim, lacunaritate, nr_boxes = calc_lacunarity_and_fractal_dim(os.path.join(r_path, file))
            for index, key in enumerate(lacunaritate):
                lacunaritate_medie += lacunaritate[key]
                # lacunaritate_tabel_terminal_header.append(str(key))
                lacunaritate_R.append(str(round(lacunaritate[key], 4)))
            # lacunaritate_tabel_terminal = PrettyTable(set(lacunaritate_tabel_terminal_header))
            lacunaritate_medie = lacunaritate_medie / len(lacunaritate)
            lacunaritate_medie_list.append(lacunaritate_medie)
            first_rows.append([filename, os.path.join(r_path, file), 'R', Ns, round(procent_box, 4), round(fractal_dim, 5), round(lacunaritate_medie, 4)])
            lacunaritate_rows.append(lacunaritate_R)
            lacunaritate_tabel_terminal.add_row(lacunaritate_R)
            dimensiune_fractala_medie_list.append(fractal_dim)
            dim_medii["r"] = fractal_dim
            lac_medii["r"] = lacunaritate_medie

            lacunaritate_G = [filename, 'G']
            Ns, procent_box, factor_divizare, fractal_dim, lacunaritate, nr_boxes = calc_lacunarity_and_fractal_dim(os.path.join(g_path, file))
            for index, key in enumerate(lacunaritate):
                lacunaritate_medie += lacunaritate[key]
                lacunaritate_G.append(str(round(lacunaritate[key], 4)))
            lacunaritate_medie = lacunaritate_medie / len(lacunaritate)
            lacunaritate_medie_list.append(lacunaritate_medie)
            first_rows.append([filename, os.path.join(g_path, file), 'G', Ns, round(procent_box, 4), round(fractal_dim, 5), round(lacunaritate_medie, 4)])
            lacunaritate_rows.append(lacunaritate_G)
            lacunaritate_tabel_terminal.add_row(lacunaritate_G)
            dimensiune_fractala_medie_list.append(fractal_dim)
            dim_medii["g"] = fractal_dim
            lac_medii["g"] = lacunaritate_medie
     
            lacunaritate_B = [filename, 'B']
            Ns, procent_box, factor_divizare, fractal_dim, lacunaritate, nr_boxes = calc_lacunarity_and_fractal_dim(os.path.join(b_path, file))
            for index, key in enumerate(lacunaritate):
                lacunaritate_medie += lacunaritate[key]
                lacunaritate_B.append(str(round(lacunaritate[key], 4)))
            lacunaritate_medie = lacunaritate_medie / len(lacunaritate)
            lacunaritate_medie_list.append(lacunaritate_medie)
            first_rows.append([filename, os.path.join(b_path, file), 'B', Ns, round(procent_box, 4), round(fractal_dim, 5), round(lacunaritate_medie, 4)])
            lacunaritate_rows.append(lacunaritate_B)
            lacunaritate_tabel_terminal.add_row(lacunaritate_B)
            dimensiune_fractala_medie_list.append(fractal_dim)
            dim_medii["b"] = fractal_dim
            lac_medii["b"] = lacunaritate_medie

            lacunaritate_tabel_terminal.add_row(['-'] * len(lacunaritate_header))
            dim_fractala_medie_rows.append([filename, os.path.join(original_images_path, original_filename), sum(dimensiune_fractala_medie_list) / len(dimensiune_fractala_medie_list), sum(lacunaritate_medie_list) / len(lacunaritate_medie_list)])

            dim_medii["medie"] = sum(dimensiune_fractala_medie_list) / len(dimensiune_fractala_medie_list)
            lac_medii["medie"] = sum(lacunaritate_medie_list) / len(lacunaritate_medie_list)

            dim_fractal_graph.append((filename, dim_medii))
            lac_graph.append((filename, lac_medii))

            sys.stdout.flush()
            sys.stdout.write("\bCurrent progress: %s %%\r" % (str(math.ceil(idx / len(files) * 100))))
    

    build_graphs(dim_fractal_graph, 'Dimensiunile fractale')
    build_graphs(lac_graph, 'Valorile lacunaritatii')
    
    build_word_doc(first_rows, first_headers, 'tabel_principal.docx')
    build_word_doc(lacunaritate_rows, lacunaritate_header, 'lacunaritate.docx')
    build_word_doc(original_images_rows, original_images_header, 'original_images.docx')
    build_word_doc(dim_fractala_medie_rows, dim_fractala_medie_header, 'dim_fractala_medie.docx')

    # lacunaritate_tabel_terminal = PrettyTable(lacunaritate_header)
    # for row in lacunaritate_rows:
    #     lacunaritate_tabel_terminal.add_row(row)
    #     lacunaritate_tabel_terminal.add_row(['-'] * len(row))
        
    sys.stdout.write("\b\nDone\n")
    print(lacunaritate_tabel_terminal)


def calc_lacunarity_and_fractal_dim(path):
    
    image = Image.open(path) # Brodatz/D1.gif
    image = image.convert('L')
    (imM, _) = image.size

    # calculate Nr and r
    Nr = []
    r = []
    
    a = 2
    b = imM//2
    nval = 20
    lnsp = np.linspace(1,math.log(b,a),nval)
    sval  = a**lnsp
    L = {}
    for S in sval:#range(2,imM//2,(imM//2-2)//100):
        Ns, N, nr = dbc(image, int(S))
        Nr.append(Ns)
        R = S/imM
        r.append(S)
        temp = 0
        L[int(S)] = 0
        N = int(math.sqrt(N))
        for i in range(N):
            for j in range(N):
                count = 0
                for k in range(N):
                    for l in range(N):
                        if nr[k][l] == nr[i][j]:
                            count = count + 1
                prob = count / (N*N)
                L[int(S)] = L[int(S)] + (i*j)**2 * prob
                temp = temp + i * prob
        L[int(S)] = L[int(S)] / (temp**2)
        # t.add_row([img_name, Ns, R, S, L[int(S)], N*N])

    # calculate log(Nr) and log(1/r)    
    y = np.log(np.array(Nr))
    x = np.log(1/np.array(r))
    (D, b) = np.polyfit(x, y, deg=1)

    # search fit error value
    N = len(x)
    Sum = 0
    for i in range(N):
        Sum += (D*x[i] + b - y[i])**2
        
    errorfit = (1/N)*math.sqrt(Sum/(1+D**2))

    # figure size 10x5 inches
    # plt.figure(1,figsize=(10,5))
    # plt.subplots_adjust(left=0.04,right=0.98)
    # plt.subplot(121)
    # plt.title(path)
    # plt.imshow(image)
    # plt.axis('off')


    # plt.subplot(122)  
    # plt.title('Fractal dimension = %f\n Fit Error = %f' % (D,errorfit))

    # plt.plot(x, y, 'ro',label='Calculated points')
    # plt.plot(x, D*x+b, 'k--', label='Linear fit' )
    # plt.legend(loc=4)
    # plt.xlabel('log(1/r)')
    # plt.ylabel('log(Nr)')
    # plt.show(block=False)
    # plt.pause(0.1)
    # plt.close()

    return Ns, R, S, D, L, N*N